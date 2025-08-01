# File: backend/parsers/text_extractor.py

import PyPDF2
import docx
import os

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file using PyPDF2. Returns clean text and page count."""
    text = ""
    page_count = 0
    try:
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            page_count = len(reader.pages)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}", 0
    return text.strip(), page_count

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file using python-docx. Returns clean text and paragraph count."""
    try:
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        para_count = len(doc.paragraphs)
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}", 0
    return text.strip(), para_count

def extract_text_and_meta(filepath):
    """
    Universal extractor: returns text and meta info (type, length, preview, file size, success flag).
    """
    text = ""
    meta = {}
    success = True
    if filepath.lower().endswith(".pdf"):
        text, count = extract_text_from_pdf(filepath)
        meta = {"type": "pdf", "pages": count}
    elif filepath.lower().endswith(".docx"):
        text, count = extract_text_from_docx(filepath)
        meta = {"type": "docx", "paragraphs": count}
    elif filepath.lower().endswith(".txt"):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()
            meta = {"type": "txt", "lines": len(text.splitlines())}
        except Exception as e:
            return f"Error extracting text from TXT: {str(e)}", {"success": False}
    else:
        return "Unsupported file format", {"success": False}

    # Add preview and file size for uniqueness
    preview = text[:200] + "..." if len(text) > 200 else text
    file_size = os.path.getsize(filepath) if os.path.exists(filepath) else 0
    meta["preview"] = preview
    meta["file_size_bytes"] = file_size
    meta["success"] = success and bool(text and len(text) > 30)  # flag if text is too short

    return text.strip(), meta
